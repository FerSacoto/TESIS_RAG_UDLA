"""
╔══════════════════════════════════════════════════════════════════╗
║       SISTEMA RAG - GENERADOR DE SÍLABOS UDLA                  ║
║       Ejecutar: python generar_silabo.py                        ║
╚══════════════════════════════════════════════════════════════════╝
"""

import os, sys, json, datetime, warnings
warnings.filterwarnings("ignore")

# ══════════════════════════════════════════════════════════════════
#  CONFIGURACIÓN — EDITA SOLO ESTA SECCIÓN
# ══════════════════════════════════════════════════════════════════

OPENAI_API_KEY = "AIzaSyC8yAoM1LrlM_-rSm2ljwpE42z3WsM6X90"   # ← REEMPLAZA con tu clave
MODEL_NAME     = "gpt-3.5-turbo"            # ← o "gpt-4o"

DATOS_SILABO = {
    "maestria":       "Maestría en Asesoría y Gerencia Legal de Empresas",
    "asignatura":     "ESCRIBE AQUÍ EL NOMBRE DE LA ASIGNATURA",   # ← EDITAR
    "area":           "ESCRIBE AQUÍ EL ÁREA TEMÁTICA",             # ← EDITAR
    "creditos":       "3",
    "sesiones":       "8",
    "horas_totales":  "48",
    "h_docente":      "16",
    "h_practico":     "8",
    "h_autonomo":     "24",
    "nombre_docente": "NOMBRE DEL DOCENTE",                        # ← EDITAR
    "perfil_docente": "PERFIL ACADÉMICO DEL DOCENTE",              # ← EDITAR
    "email_docente":  "correo@udlaonline.edu.ec",                  # ← EDITAR
}

# ══════════════════════════════════════════════════════════════════
#  RUTAS — adaptadas a tu estructura de carpetas
# ══════════════════════════════════════════════════════════════════

from pathlib import Path

BASE      = Path(__file__).parent.resolve()
DATA      = BASE / "data"
OUTPUT    = BASE / "output"
TEMPLATES = BASE / "templates"
INDICES   = DATA / "indices"

# Tus archivos reales
PDF_MODELO_EDUCATIVO = DATA / "Modelo-Educativo-UDLA-2025-2030.pdf"
EXCEL_ADN            = DATA / "Tabla ADN_UDLA5H111_Maestrpua en Asesoría y Gerencia Legal de Empresas.xlsx"
EXCEL_MULTIMEDIA     = DATA / "dataset_recursosmultimedia.xlsx"
EXCEL_BIBLIOGRAFICA  = DATA / "dataset_referenciasbiliográficas.xlsx"
PLANTILLA            = TEMPLATES / "DEL_02_PLANTILLA_SILABO_UDLA_LIMPIA.docx"

# Índices FAISS
IDX_NORMATIVA    = INDICES / "faiss_normativa"
IDX_MULTIMEDIA   = INDICES / "faiss_multimedia"
IDX_BIBLIOGRAFICA= INDICES / "faiss_bibliografica"

# Parámetros RAG
TOP_K         = 5
CHUNK_SIZE    = 800
CHUNK_OVERLAP = 150
EMBED_MODEL   = "text-embedding-3-small"

# ══════════════════════════════════════════════════════════════════
#  IMPORTS
# ══════════════════════════════════════════════════════════════════

def verificar_imports():
    faltantes = []
    for mod, pkg in [
        ("langchain","langchain"),("langchain_community","langchain-community"),
        ("langchain_openai","langchain-openai"),("faiss","faiss-cpu"),
        ("pypdf","pypdf"),("openpyxl","openpyxl"),
        ("docx","python-docx"),("openai","openai"),
        ("pandas","pandas"),("tiktoken","tiktoken"),
    ]:
        try: __import__(mod)
        except ImportError: faltantes.append(pkg)
    if faltantes:
        print(f"\n❌ Falta instalar: pip install {' '.join(faltantes)}\n")
        sys.exit(1)

verificar_imports()

import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.schema import Document
from langchain.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from docx import Document as DocxDocument
from docx.shared import Pt

# ══════════════════════════════════════════════════════════════════
#  CARGA DE ARCHIVOS
# ══════════════════════════════════════════════════════════════════

def cargar_pdf(ruta: Path) -> list:
    if not ruta.exists():
        print(f"   ⚠  No encontrado: {ruta.name}")
        return []
    print(f"   📄 {ruta.name}")
    try:
        loader = PyPDFLoader(str(ruta))
        docs = loader.load()
        for d in docs:
            d.metadata["fuente"] = ruta.name
        print(f"      → {len(docs)} páginas")
        return docs
    except Exception as e:
        print(f"      ⚠  Error: {e}")
        return []


def cargar_excel(ruta: Path, tipo: str) -> list:
    if not ruta.exists():
        print(f"   ⚠  No encontrado: {ruta.name}")
        return []
    print(f"   📊 {ruta.name}")
    try:
        xls  = pd.ExcelFile(str(ruta))
        docs = []
        for hoja in xls.sheet_names:
            df = xls.parse(hoja).fillna("")
            for _, fila in df.iterrows():
                texto = "\n".join(
                    f"{col}: {val}" for col, val in fila.items() if str(val).strip()
                )
                if texto.strip():
                    docs.append(Document(
                        page_content=texto,
                        metadata={"fuente": ruta.name, "hoja": hoja, "tipo": tipo}
                    ))
        print(f"      → {len(docs)} filas indexadas")
        return docs
    except Exception as e:
        print(f"      ⚠  Error: {e}")
        return []


def construir_o_cargar(docs, ruta_idx: Path, nombre: str, emb):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    if ruta_idx.exists():
        print(f"   ♻  Cargando índice existente: {nombre}")
        return FAISS.load_local(str(ruta_idx), emb,
                                allow_dangerous_deserialization=True)
    if not docs:
        print(f"   ⚠  Sin documentos para: {nombre}")
        return None
    print(f"   🔧 Construyendo índice FAISS: {nombre}")
    fragmentos = splitter.split_documents(docs)
    print(f"      → {len(fragmentos)} fragmentos")
    idx = FAISS.from_documents(fragmentos, emb)
    ruta_idx.mkdir(parents=True, exist_ok=True)
    idx.save_local(str(ruta_idx))
    print(f"      ✅ Guardado en data/indices/{ruta_idx.name}/")
    return idx


def recuperar(idx, query: str, etiqueta: str) -> str:
    if idx is None:
        return f"[{etiqueta}: no disponible]"
    docs = idx.as_retriever(search_kwargs={"k": TOP_K}).invoke(query)
    if not docs:
        return f"[{etiqueta}: sin resultados]"
    return "\n\n".join(
        f"[{etiqueta} #{i} | {d.metadata.get('fuente','')}]\n{d.page_content}"
        for i, d in enumerate(docs, 1)
    )

# ══════════════════════════════════════════════════════════════════
#  PROMPTS
# ══════════════════════════════════════════════════════════════════

P_DESCRIPCION = """
Eres experto en diseño curricular universitario de posgrado.
Redacta la DESCRIPCIÓN DEL CURSO para la asignatura indicada.
Debe ser clara, académica, entre 150 y 250 palabras, en español formal,
alineada con el Modelo Educativo UDLA 2025-2030.

=== MODELO EDUCATIVO Y ADN CURRICULAR ===
{normativa}

=== DATOS DEL CURSO ===
Maestría: {maestria}
Asignatura: {asignatura}
Área temática: {area}
Créditos: {creditos}

Redacta únicamente la descripción:
"""

P_RDA = """
Eres experto en taxonomía de Bloom y resultados de aprendizaje de posgrado.
Genera entre 3 y 4 RESULTADOS DE APRENDIZAJE (RdA) para la asignatura.
Cada RdA inicia con verbo de acción de Bloom en infinitivo, es medible
y está alineado al ADN curricular del programa.

=== MODELO EDUCATIVO Y ADN CURRICULAR ===
{normativa}

=== DATOS DEL CURSO ===
Maestría: {maestria}
Asignatura: {asignatura}
Área temática: {area}

Genera los RdAs numerados (RdA 1, RdA 2, RdA 3...):
"""

P_CRONOGRAMA = """
Eres experto en planificación didáctica universitaria de posgrado.
Genera el cronograma semanal en formato JSON puro (sin texto extra ni markdown).
Incluye los recursos multimedia disponibles como actividades concretas.

=== MODELO EDUCATIVO ===
{normativa}

=== RECURSOS MULTIMEDIA DISPONIBLES ===
{multimedia}

=== DATOS DEL CURSO ===
Maestría: {maestria}
Asignatura: {asignatura}
Número de sesiones: {sesiones}
Resultados de Aprendizaje: {rda}

Devuelve SOLO el array JSON con esta estructura (una entrada por semana):
[
  {{
    "semana": "Semana 1",
    "sesion": "1",
    "rda1": "✓",
    "rda2": "",
    "rda3": "",
    "pre_sesion": "Actividad pre-sesión...",
    "sincronica": "Contenido sesión sincrónica...",
    "post_sesion": "Actividad post-sesión...",
    "evaluacion": "Nombre o vacío",
    "ponderacion": "% o vacío"
  }}
]
"""

P_REFERENCIAS = """
Eres bibliotecólogo académico especializado en APA 7ma edición.
Selecciona del catálogo las referencias más pertinentes para la asignatura.
Organiza en PRINCIPALES (3-5) y COMPLEMENTARIAS (3-4).

=== CATÁLOGO BIBLIOGRÁFICO INSTITUCIONAL ===
{bibliografica}

=== DATOS DEL CURSO ===
Asignatura: {asignatura}
Área temática: {area}

Formato APA 7. Responde así:
PRINCIPALES:
- Autor, A. (año). Título. Editorial.

COMPLEMENTARIAS:
- Autor, A. (año). Título. Editorial.
"""

# ══════════════════════════════════════════════════════════════════
#  GENERACIÓN CON LLM
# ══════════════════════════════════════════════════════════════════

def generar(llm, prompt: str, variables: dict) -> str:
    chain = ChatPromptTemplate.from_template(prompt) | llm | StrOutputParser()
    return chain.invoke(variables)

# ══════════════════════════════════════════════════════════════════
#  RELLENO DE PLANTILLA DOCX
# ══════════════════════════════════════════════════════════════════

def reemplazar_parrafo(parrafo, mapa: dict):
    texto = parrafo.text
    for ph, val in mapa.items():
        if ph in texto:
            for run in parrafo.runs:
                run.text = ""
            if parrafo.runs:
                parrafo.runs[0].text = texto.replace(ph, str(val))
            else:
                parrafo.add_run(texto.replace(ph, str(val)))
            return


def rellenar_cronograma(doc, cronograma: list):
    CAMPOS = ["semana","sesion","rda1","rda2","rda3",
              "pre_sesion","sincronica","post_sesion","evaluacion","ponderacion"]
    for tabla in doc.tables:
        if tabla.rows and "Semana" in tabla.rows[0].cells[0].text:
            while len(tabla.rows) > 1:
                tabla._tbl.remove(tabla.rows[-1]._tr)
            for entrada in cronograma:
                fila = tabla.add_row()
                for i, campo in enumerate(CAMPOS):
                    if i < len(fila.cells):
                        p = fila.cells[i].paragraphs[0]
                        p.clear()
                        p.add_run(str(entrada.get(campo, ""))).font.size = Pt(8)
            print(f"   ✅ {len(cronograma)} semanas insertadas en la tabla")
            return
    print("   ⚠  No se encontró la tabla del cronograma")


def llenar_plantilla(datos: dict, cronograma: list) -> Path:
    if not PLANTILLA.exists():
        print(f"❌ Plantilla no encontrada en: {PLANTILLA}")
        sys.exit(1)

    doc = DocxDocument(str(PLANTILLA))

    MAPA = {
        "{{ MAESTRIA }}":            datos.get("maestria",""),
        "{{ ASIGNATURA }}":          datos.get("asignatura",""),
        "{{ creditos }}":            datos.get("creditos",""),
        "{{ sesiones }}":            datos.get("sesiones",""),
        "{{ horas_totales }}":       datos.get("horas_totales",""),
        "{{ h_docente }}":           datos.get("h_docente",""),
        "{{ h_practico }}":          datos.get("h_practico",""),
        "{{ h_autonomo }}":          datos.get("h_autonomo",""),
        "{{ descripcion_curso }}":   datos.get("descripcion_curso",""),
        "{{ bloque_rda }}":          datos.get("bloque_rda",""),
        "{{ nombre_docente }}":      datos.get("nombre_docente",""),
        "{{ perfil_docente }}":      datos.get("perfil_docente",""),
        "{{ email_docente }}":       datos.get("email_docente",""),
        "{{ ref_principales }}":     datos.get("ref_principales",""),
        "{{ ref_complementarias }}": datos.get("ref_complementarias",""),
    }

    for p in doc.paragraphs:
        reemplazar_parrafo(p, MAPA)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    reemplazar_parrafo(p, MAPA)

    if cronograma:
        rellenar_cronograma(doc, cronograma)

    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    nom  = f"SILABO_{datos['asignatura'][:30].replace(' ','_')}_{ts}.docx"
    ruta = OUTPUT / nom
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc.save(str(ruta))
    return ruta

# ══════════════════════════════════════════════════════════════════
#  PROGRAMA PRINCIPAL
# ══════════════════════════════════════════════════════════════════

def main():
    print("\n" + "═"*60)
    print("  🎓  SISTEMA RAG — GENERADOR DE SÍLABOS UDLA")
    print("═"*60)

    if OPENAI_API_KEY.startswith("sk-PEGA"):
        print("\n❌  Edita generar_silabo.py y pega tu OpenAI API Key")
        print("   Busca la línea:  OPENAI_API_KEY = \"sk-PEGA_AQUI_TU_CLAVE\"")
        sys.exit(1)

    if "ESCRIBE AQUÍ" in DATOS_SILABO["asignatura"]:
        print("\n❌  Edita generar_silabo.py y completa DATOS_SILABO")
        print("   Busca la sección:  DATOS_SILABO = { ... }")
        sys.exit(1)

    os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY
    INDICES.mkdir(parents=True, exist_ok=True)

    print("\n⚙  Iniciando modelos de IA...")
    emb = OpenAIEmbeddings(model=EMBED_MODEL)
    llm = ChatOpenAI(model=MODEL_NAME, temperature=0.3)
    print(f"   Modelo : {MODEL_NAME}")

    # ── Cargar BD NORMATIVA ──────────────────────────────────────
    print("\n📋 [1/3] Cargando BD NORMATIVA...")
    docs_norm  = cargar_pdf(PDF_MODELO_EDUCATIVO)
    docs_norm += cargar_excel(EXCEL_ADN, "adn_curricular")
    idx_norm   = construir_o_cargar(docs_norm, IDX_NORMATIVA, "Normativa", emb)

    # ── Cargar BD MULTIMEDIA ─────────────────────────────────────
    print("\n🎬 [2/3] Cargando BD MULTIMEDIA...")
    docs_mm = cargar_excel(EXCEL_MULTIMEDIA, "recurso_multimedia")
    idx_mm  = construir_o_cargar(docs_mm, IDX_MULTIMEDIA, "Multimedia", emb)

    # ── Cargar BD BIBLIOGRÁFICA ──────────────────────────────────
    print("\n📖 [3/3] Cargando BD BIBLIOGRÁFICA...")
    docs_bib = cargar_excel(EXCEL_BIBLIOGRAFICA, "referencia_bibliografica")
    idx_bib  = construir_o_cargar(docs_bib, IDX_BIBLIOGRAFICA, "Bibliográfica", emb)

    # ── Recuperar contexto ───────────────────────────────────────
    print("\n🔍 Recuperando contexto de las 3 bases de datos...")
    query    = f"{DATOS_SILABO['asignatura']} {DATOS_SILABO['area']}"
    ctx_norm = recuperar(idx_norm, query, "NORMATIVA")
    ctx_mm   = recuperar(idx_mm,   query, "MULTIMEDIA")
    ctx_bib  = recuperar(idx_bib,  query, "BIBLIOGRAFÍA")
    print("   ✅ Contexto recuperado")

    resultado = {**DATOS_SILABO}

    # ── Generar secciones ────────────────────────────────────────
    print("\n✍  Generando descripción del curso...")
    resultado["descripcion_curso"] = generar(llm, P_DESCRIPCION,
        {**DATOS_SILABO, "normativa": ctx_norm})
    print("   ✅ Listo")

    print("\n🎯 Generando Resultados de Aprendizaje...")
    resultado["bloque_rda"] = generar(llm, P_RDA,
        {**DATOS_SILABO, "normativa": ctx_norm})
    print("   ✅ Listo")

    print("\n📅 Generando cronograma semanal...")
    cron_raw = generar(llm, P_CRONOGRAMA, {
        **DATOS_SILABO,
        "normativa":  ctx_norm,
        "multimedia": ctx_mm,
        "rda":        resultado["bloque_rda"],
    })
    cron_raw = cron_raw.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
    try:
        cronograma = json.loads(cron_raw)
        print(f"   ✅ {len(cronograma)} semanas generadas")
    except json.JSONDecodeError:
        print("   ⚠  Error al leer JSON del cronograma. Se dejará vacío.")
        cronograma = []

    print("\n📚 Seleccionando referencias bibliográficas...")
    refs = generar(llm, P_REFERENCIAS,
        {**DATOS_SILABO, "bibliografica": ctx_bib})
    if "COMPLEMENTARIAS:" in refs:
        partes = refs.split("COMPLEMENTARIAS:")
        resultado["ref_principales"]     = partes[0].replace("PRINCIPALES:","").strip()
        resultado["ref_complementarias"] = partes[1].strip()
    else:
        resultado["ref_principales"]     = refs.strip()
        resultado["ref_complementarias"] = "(Ver catálogo de biblioteca institucional)"
    print("   ✅ Listo")

    print("\n📄 Generando documento Word final...")
    ruta_final = llenar_plantilla(resultado, cronograma)

    print("\n" + "═"*60)
    print("  🎉  ¡SÍLABO GENERADO EXITOSAMENTE!")
    print("═"*60)
    print(f"\n  📁 Archivo guardado en:")
    print(f"     output\\{ruta_final.name}\n")


if __name__ == "__main__":
    main()
